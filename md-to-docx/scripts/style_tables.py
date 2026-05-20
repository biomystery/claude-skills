#!/usr/bin/env python3
"""Apply dark-navy header row and alternating row bands to all tables in a .docx file.

Usage: style_tables.py <out_docx>
Modifies the file in place.
"""
import sys
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HEADER_BG = '1E3A5F'   # dark navy — matches heading colour
ODD_BG    = 'EBF4FD'   # light blue — matches blockquote background
EVEN_BG   = 'FFFFFF'

def set_cell_shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for el in tcPr.findall(qn('w:shd')):
        tcPr.remove(el)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

doc = Document(sys.argv[1])
for table in doc.tables:
    seen = set()
    for row_idx, row in enumerate(table.rows):
        is_header = row_idx == 0
        fill = HEADER_BG if is_header else (ODD_BG if row_idx % 2 == 1 else EVEN_BG)
        for cell in row.cells:
            cid = id(cell._tc)
            if cid in seen:
                continue
            seen.add(cid)
            set_cell_shading(cell, fill)
            if is_header:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
doc.save(sys.argv[1])
